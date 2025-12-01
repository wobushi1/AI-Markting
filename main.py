import sys
import json
import base64
import os
from io import BytesIO
from typing import List

# PyQt5 界面库
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                             QHBoxLayout, QPushButton, QTextEdit, QLabel, 
                             QFileDialog, QListWidget, QSplitter, QProgressBar,
                             QLineEdit, QFormLayout, QMessageBox, QTabWidget)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QFont, QColor

# 图片与PDF处理
from PIL import Image
from pdf2image import convert_from_path

# AI 模型接口
from openai import OpenAI

# Word 导出库
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn  # 关键：修复字体设置报错

# ==========================================
# 工具函数：处理 Poppler 路径 (跨平台打包关键)
# ==========================================
def get_poppler_path():
    """
    智能获取 Poppler 路径：
    1. Linux/Mac: 返回 None (假设系统已安装 poppler-utils)
    2. Windows (打包后): 返回临时目录下的 poppler/bin
    3. Windows (开发中): 返回当前目录下的 poppler/bin
    """
    if sys.platform != "win32":
        return None  # Linux 环境通常不需要指定路径，只要安装了 poppler-utils
    
    # 获取基础路径：如果是 exe 运行则是临时目录 _MEIPASS，否则是当前脚本目录
    base_path = getattr(sys, '_MEIPASS', os.path.abspath("."))
    
    # 构造 poppler/bin 的绝对路径
    # 注意：这要求打包时使用了 --add-data "poppler;poppler"
    return os.path.join(base_path, 'poppler', 'bin')

# ==========================================
# 评分标准 Prompt (JSON 结构化输出)
# ==========================================
RUBRIC_PROMPT = """
你是一位资深的高考英语阅卷专家。请对上传的手写英语作文图片进行识别、分类、评分，并提供极度详细的逐句修改意见。

### 任务步骤：
1. **OCR识别**：准确识别图片内容。
2. **分类**：应用文 vs 读后续写。
3. **评分**：基于高考标准（满分15分）打分。
4. **深度反馈（关键）**：
   - **内容要点**：必须分开列出“不足”和“建议”。
   - **语言表达**：**必须**提取文中3-5个典型错误或需提升的句子，进行逐句修改，并解释语法点/词汇选择原因。
   - **结构**：评价逻辑衔接。
5. **全文润色**：提供一篇满分范文。

### 输出格式要求：
请**务必**仅返回纯合法的 JSON 格式字符串，不要包含 Markdown 标记（如 ```json），JSON结构严格如下：

{
    "recognized_text": "识别出的原文...",
    "essay_type": "应用文",
    "scores": {
        "dim1_score": 4, 
        "dim2_score": 3,
        "dim3_score": 4,
        "total": 11
    },
    "feedback_detail": {
        "content": {
            "weakness": "...",
            "suggestion": "..."
        },
        "language": {
            "sentence_corrections": [
                {
                    "original": "I can relate you pain of it.",
                    "revised": "I can relate to your pain.",
                    "explanation": "‘relate to’是固定搭配..."
                }
            ],
            "general_comment": "整体语言风格评价..."
        },
        "structure": "评价文章的逻辑结构...",
        "overall_summary": "整体优缺点总结..."
    },
    "revised_version": "Full revised essay..."
}
"""

# ==========================================
# 后端工作线程
# ==========================================
class Worker(QThread):
    finished = pyqtSignal(dict, str) # 返回结果JSON和文件路径
    error = pyqtSignal(str, str)     # 返回错误信息和文件路径

    def __init__(self, file_path, api_key, model_endpoint):
        super().__init__()
        self.file_path = file_path
        self.api_key = api_key
        self.model_endpoint = model_endpoint 

    def encode_image(self, image_path):
        """处理图片/PDF转Base64"""
        img = None
        ext = os.path.splitext(image_path)[1].lower()

        try:
            # 1. 处理 PDF
            if ext == '.pdf':
                try:
                    # 获取适配当前系统的 poppler 路径
                    poppler_dir = get_poppler_path()
                    
                    pages = convert_from_path(
                        image_path, 
                        first_page=1, 
                        last_page=1, 
                        poppler_path=poppler_dir # 传入路径
                    )
                    if pages: img = pages[0]
                except Exception as e:
                    raise Exception(f"PDF处理失败: {str(e)}\n(如果是Windows，请检查Poppler路径配置)")
            
            # 2. 处理图片
            else:
                img = Image.open(image_path)

            if img is None: raise Exception("无法加载文件内容")

            # 3. 格式统一化
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            
            # 缩放限制
            max_size = 2048
            if max(img.size) > max_size:
                img.thumbnail((max_size, max_size), Image.Resampling.LANCZOS)

            buffered = BytesIO()
            img.save(buffered, format="JPEG", quality=85)
            return base64.b64encode(buffered.getvalue()).decode('utf-8')

        except Exception as e:
            raise Exception(f"图片预处理失败: {str(e)}")

    def run(self):
        try:
            base64_image = self.encode_image(self.file_path)
            
            client = OpenAI(
                api_key=self.api_key,
                base_url="https://ark.cn-beijing.volces.com/api/v3", 
            )

            response = client.chat.completions.create(
                model=self.model_endpoint,
                messages=[
                    {"role": "system", "content": RUBRIC_PROMPT},
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": "请批改这张作文图片，请严格按照JSON格式返回。"},
                            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                        ]
                    }
                ],
                temperature=0.2 # 低随机性，保证JSON格式稳定
            )

            content = response.choices[0].message.content
            # 清理 Markdown 标记
            content = content.replace("```json", "").replace("```", "").strip()
            
            try:
                result_json = json.loads(content)
                self.finished.emit(result_json, self.file_path)
            except json.JSONDecodeError:
                self.error.emit(f"AI返回格式异常，无法解析 JSON。\n原始内容片段:\n{content[:200]}", self.file_path)

        except Exception as e:
            self.error.emit(str(e), self.file_path)

# ==========================================
# 前端 GUI
# ==========================================
class EssayGraderApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("高考英语作文智能批改系统 (Pro版)")
        self.resize(1280, 850)
        
        # 核心：存储所有文件的结果 {filepath: json_data}
        self.results_store = {}
        
        self.init_ui()

    def init_ui(self):
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        layout = QVBoxLayout(main_widget)

        # 1. 配置区
        config_layout = QFormLayout()
        self.api_key_input = QLineEdit()
        self.api_key_input.setPlaceholderText("火山引擎 API Key")
        self.api_key_input.setEchoMode(QLineEdit.Password)
        self.endpoint_input = QLineEdit()
        self.endpoint_input.setPlaceholderText("接入点 ID (如 ep-2024... Vision版)")
        config_layout.addRow("API Key:", self.api_key_input)
        config_layout.addRow("Endpoint ID:", self.endpoint_input)
        layout.addLayout(config_layout)

        # 2. 主区域
        splitter = QSplitter(Qt.Horizontal)
        
        # --- 左侧控制区 ---
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        
        btn_layout = QHBoxLayout()
        self.btn_add = QPushButton("添加文件")
        self.btn_add.clicked.connect(self.add_files)
        
        self.btn_run = QPushButton("开始批改")
        self.btn_run.clicked.connect(self.start_grading)
        self.btn_run.setStyleSheet("background-color: #007AFF; color: white; font-weight: bold;")
        
        btn_layout.addWidget(self.btn_add)
        btn_layout.addWidget(self.btn_run)
        
        # 导出按钮
        self.btn_export = QPushButton("导出Word报告")
        self.btn_export.clicked.connect(self.export_to_word)
        self.btn_export.setStyleSheet("background-color: #FF9800; color: white; font-weight: bold;")
        self.btn_export.setEnabled(False)

        left_layout.addLayout(btn_layout)
        left_layout.addWidget(self.btn_export)
        left_layout.addWidget(QLabel("文件列表 (点击查看详情):"))
        
        self.file_list = QListWidget()
        self.file_list.itemClicked.connect(self.load_selected_result) # 绑定点击事件
        left_layout.addWidget(self.file_list)
        
        # --- 右侧 Tab 显示区 ---
        right_widget = QTabWidget()
        right_widget.setStyleSheet("QTextEdit { font-size: 14px; line-height: 1.6; }")
        
        self.text_original = QTextEdit()
        self.text_original.setReadOnly(True)
        right_widget.addTab(self.text_original, "📝 识别原文")
        
        self.text_feedback = QTextEdit()
        self.text_feedback.setReadOnly(True)
        right_widget.addTab(self.text_feedback, "📊 深度精批")
        
        self.text_revised = QTextEdit()
        self.text_revised.setReadOnly(True)
        right_widget.addTab(self.text_revised, "✨ 满分范文")

        splitter.addWidget(left_widget)
        splitter.addWidget(right_widget)
        splitter.setStretchFactor(1, 3) 
        layout.addWidget(splitter)

        # 3. 状态栏
        self.progress_bar = QProgressBar()
        self.status_label = QLabel("就绪")
        layout.addWidget(self.progress_bar)
        layout.addWidget(self.status_label)

        self.setFont(QFont("Microsoft YaHei", 10))

    def add_files(self):
        filters = "All Support (*.png *.jpg *.jpeg *.pdf);;Images (*.png *.jpg);;PDF (*.pdf)"
        files, _ = QFileDialog.getOpenFileNames(self, "选择文件", "", filters)
        if files:
            for f in files:
                # 检查是否已存在
                items = [self.file_list.item(x).data(Qt.UserRole) for x in range(self.file_list.count())]
                if f not in items:
                    item_name = os.path.basename(f)
                    list_item = self.file_list.addItem(item_name)
                    # 将完整路径存入 item 数据
                    self.file_list.item(self.file_list.count()-1).setData(Qt.UserRole, f)

    def start_grading(self):
        if self.file_list.count() == 0: return
        api_key = self.api_key_input.text().strip()
        endpoint = self.endpoint_input.text().strip()
        if not api_key or not endpoint:
            QMessageBox.warning(self, "提示", "请填写API Key和Endpoint")
            return
        
        self.btn_run.setEnabled(False)
        self.btn_export.setEnabled(False)
        self.process_next_file(0, api_key, endpoint)

    def process_next_file(self, index, api_key, endpoint):
        if index >= self.file_list.count():
            self.status_label.setText("所有文件处理完成")
            self.progress_bar.setValue(100)
            self.btn_run.setEnabled(True)
            self.btn_export.setEnabled(True)
            QMessageBox.information(self, "完成", "批改完成，现在可以导出Word了。")
            return

        item = self.file_list.item(index)
        file_path = item.data(Qt.UserRole)
        
        # 如果已批改过，跳过
        if file_path in self.results_store:
            self.process_next_file(index + 1, api_key, endpoint)
            return

        self.file_list.setCurrentRow(index)
        self.status_label.setText(f"正在处理: {os.path.basename(file_path)}")
        self.progress_bar.setValue(int((index / self.file_list.count()) * 100))

        self.worker = Worker(file_path, api_key, endpoint)
        self.worker.finished.connect(lambda res, path: self.on_result(res, path, index, api_key, endpoint))
        self.worker.error.connect(lambda err, path: self.on_error(err, path, index, api_key, endpoint))
        self.worker.start()

    def on_result(self, result, file_path, index, api_key, endpoint):
        # 存入字典
        self.results_store[file_path] = result
        
        # 更新列表状态
        self.file_list.item(index).setText(f"[√] {os.path.basename(file_path)}")
        self.file_list.item(index).setForeground(QColor("green"))
        
        # 展示当前
        self.display_result(result)
        
        # 继续下一个
        self.process_next_file(index + 1, api_key, endpoint)

    def on_error(self, err, file_path, index, api_key, endpoint):
        self.status_label.setText(f"错误: {err}")
        self.file_list.item(index).setText(f"[X] {os.path.basename(file_path)}")
        self.file_list.item(index).setForeground(QColor("red"))
        # 出错不中断，继续下一个
        self.process_next_file(index + 1, api_key, endpoint)

    def load_selected_result(self, item):
        """点击列表时回调"""
        file_path = item.data(Qt.UserRole)
        if file_path in self.results_store:
            self.display_result(self.results_store[file_path])
        else:
            self.text_original.setText("尚未批改或处理失败")
            self.text_feedback.clear()
            self.text_revised.clear()

    def display_result(self, data):
        """渲染 HTML 结果"""
        # 原文
        self.text_original.setText(f"【类型】：{data.get('essay_type')}\n\n{data.get('recognized_text')}")
        # 范文
        self.text_revised.setText(data.get('revised_version'))
        
        # 深度反馈
        scores = data.get('scores', {})
        fb = data.get('feedback_detail', {})
        content_fb = fb.get('content', {})
        lang_fb = fb.get('language', {})
        
        html = f"""
        <h2 style='color:#333'>总分：<span style='color:#E53935; font-size:24px'>{scores.get('total')}/15</span></h2>
        
        <table border='1' cellpadding='6' cellspacing='0' style='border-collapse:collapse; width:100%; border-color:#ddd;'>
            <tr style='background-color:#f5f5f5'>
                <th width='33%'>内容要点</th><th width='33%'>语言表达</th><th width='33%'>结构衔接</th>
            </tr>
            <tr>
                <td align='center'>{scores.get('dim1_score')}/5</td>
                <td align='center'>{scores.get('dim2_score')}/5</td>
                <td align='center'>{scores.get('dim3_score')}/5</td>
            </tr>
        </table>

        <h3 style='background-color:#E3F2FD; padding:5px'>一、内容要点</h3>
        <ul>
            <li><b>🔻 不足：</b> {content_fb.get('weakness', '无')}</li>
            <li><b>💡 建议：</b> {content_fb.get('suggestion', '无')}</li>
        </ul>

        <h3 style='background-color:#FFF3E0; padding:5px'>二、语言表达 (逐句精改)</h3>
        """
        
        corrections = lang_fb.get('sentence_corrections', [])
        if corrections:
            for idx, item in enumerate(corrections, 1):
                html += f"""
                <div style='margin-bottom:15px; border-bottom:1px dashed #ccc; padding-bottom:10px;'>
                    <p style='margin:4px 0'><b>{idx}. 🔴 原句：</b> <span style='color:#555'>{item.get('original')}</span></p>
                    <p style='margin:4px 0'><b>🟢 修改：</b> <span style='color:#2E7D32; font-weight:bold'>{item.get('revised')}</span></p>
                    <p style='margin:4px 0; color:#1565C0'><b>📘 解析：</b> {item.get('explanation')}</p>
                </div>
                """
        else:
            html += "<p>暂无具体句子修改建议。</p>"
            
        html += f"""
        <p><b>整体评价：</b> {lang_fb.get('general_comment', '')}</p>

        <h3 style='background-color:#E8F5E9; padding:5px'>三、结构与衔接</h3>
        <p>{fb.get('structure', '无')}</p>

        <hr>
        <p><b>🌟 整体总结：</b> {fb.get('overall_summary', '')}</p>
        """
        self.text_feedback.setHtml(html)

    # ==========================================
    # 修复后的 Word 导出逻辑
    # ==========================================
    def export_to_word(self):
        if not self.results_store:
            QMessageBox.warning(self, "提示", "没有可导出的数据")
            return

        save_path, _ = QFileDialog.getSaveFileName(self, "保存Word文档", "高考作文批改报告.docx", "Word Files (*.docx)")
        if not save_path:
            return

        doc = Document()
        
        # --- 核心：设置中文字体 ---
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei') 
        # ------------------------

        for file_path, data in self.results_store.items():
            filename = os.path.basename(file_path)
            
            # 1. 标题
            doc.add_heading(f"文件：{filename}", level=1)
            
            # 2. 原文
            doc.add_heading("OCR 识别原文", level=2)
            p = doc.add_paragraph(data.get('recognized_text', ''))
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 3. 评分表
            doc.add_heading("评分详情", level=2)
            scores = data.get('scores', {})
            table = doc.add_table(rows=2, cols=4)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = '维度', '内容要点', '语言表达', '结构衔接'
            
            row = table.rows[1].cells
            row[0].text = '得分'
            row[1].text = str(scores.get('dim1_score', 0))
            row[2].text = str(scores.get('dim2_score', 0))
            row[3].text = str(scores.get('dim3_score', 0))
            
            total_p = doc.add_paragraph()
            run = total_p.add_run(f"总分：{scores.get('total')}/15")
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.size = Pt(14)

            # 4. 深度反馈
            fb = data.get('feedback_detail', {})
            
            # 4.1 内容
            doc.add_heading("一、内容要点", level=3)
            content_fb = fb.get('content', {})
            weakness = content_fb.get('weakness') if content_fb else "无"
            suggestion = content_fb.get('suggestion') if content_fb else "无"
            doc.add_paragraph(f"不足：{weakness}", style='List Bullet')
            doc.add_paragraph(f"建议：{suggestion}", style='List Bullet')
            
            # 4.2 语言（逐句）
            doc.add_heading("二、语言表达与逐句修改", level=3)
            lang_fb = fb.get('language', {})
            corrections = lang_fb.get('sentence_corrections', []) if lang_fb else []
            
            if corrections:
                for i, item in enumerate(corrections, 1):
                    p = doc.add_paragraph()
                    p.add_run(f"{i}. 原句：").bold = True
                    p.add_run(item.get('original', '')).font.color.rgb = RGBColor(100, 100, 100)
                    
                    p = doc.add_paragraph()
                    p.add_run(f"   修改：").bold = True
                    run_rev = p.add_run(item.get('revised', ''))
                    run_rev.font.color.rgb = RGBColor(0, 128, 0)
                    run_rev.bold = True
                    
                    p = doc.add_paragraph()
                    p.add_run(f"   解析：").bold = True
                    p.add_run(item.get('explanation', '')).font.color.rgb = RGBColor(0, 0, 255)
                    doc.add_paragraph("") 
            else:
                doc.add_paragraph("暂无具体修改建议。")

            # 4.3 结构与总结
            doc.add_heading("三、结构与整体总结", level=3)
            doc.add_paragraph(f"结构评价：{fb.get('structure', '无')}")
            doc.add_paragraph(f"整体总结：{fb.get('overall_summary', '无')}")

            # 5. 范文
            doc.add_heading("满分范文参考", level=2)
            doc.add_paragraph(data.get('revised_version', '暂无'))

            doc.add_page_break()

        try:
            doc.save(save_path)
            QMessageBox.information(self, "成功", f"报告已保存至：\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "保存失败", f"错误详情：{str(e)}\n可能是文件被占用。")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = EssayGraderApp()
    window.show()
    sys.exit(app.exec_())